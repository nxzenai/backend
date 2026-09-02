from __future__ import annotations

import io
import re
import zipfile
from pathlib import Path
from typing import Any

import pandas as pd


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".xlsx", ".py", ".sql", ".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
SUPPORTED_CONTENT_TYPES = {
    ".pdf": {"application/pdf", "application/octet-stream"},
    ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/octet-stream"},
    ".txt": {"text/plain", "application/octet-stream"},
    ".csv": {"text/csv", "application/csv", "text/plain", "application/vnd.ms-excel", "application/octet-stream"},
    ".xlsx": {"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/octet-stream"},
    ".py": {"text/x-python", "text/python", "text/plain", "application/x-python-code", "application/octet-stream"},
    ".sql": {"application/sql", "application/x-sql", "text/sql", "text/plain", "application/octet-stream"},
    ".png": {"image/png", "application/octet-stream"},
    ".jpg": {"image/jpeg", "application/octet-stream"},
    ".jpeg": {"image/jpeg", "application/octet-stream"},
    ".webp": {"image/webp", "application/octet-stream"},
    ".bmp": {"image/bmp", "image/x-ms-bmp", "application/octet-stream"},
    ".tif": {"image/tiff", "application/octet-stream"},
    ".tiff": {"image/tiff", "application/octet-stream"},
}


def _clean(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value)).strip()


def validate_attachment_type(filename: str, content_type: str) -> None:
    suffix = Path(filename).suffix.casefold()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError("Supported files are PDF, DOCX, TXT, CSV, XLSX, PY, SQL, PNG, JPG, WEBP, BMP, and TIFF.")
    normalized_type = (content_type or "application/octet-stream").split(";", 1)[0].strip().casefold()
    if normalized_type not in SUPPORTED_CONTENT_TYPES[suffix]:
        raise ValueError("The file type does not match its extension.")


def extract_text(filename: str, content: bytes) -> tuple[str, dict[str, Any]]:
    suffix = Path(filename).suffix.casefold()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError("Supported files are PDF, DOCX, TXT, CSV, XLSX, PY, SQL, PNG, JPG, WEBP, BMP, and TIFF.")
    if suffix in {".docx", ".xlsx"}:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                entries = archive.infolist()
                if len(entries) > 5000 or sum(item.file_size for item in entries) > 50_000_000:
                    raise ValueError("The compressed document expands beyond the safe extraction limit.")
        except zipfile.BadZipFile as exc:
            raise ValueError("The document archive is invalid or corrupt.") from exc
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}:
        from PIL import Image

        with Image.open(io.BytesIO(content)) as image:
            image.verify()
        with Image.open(io.BytesIO(content)) as image:
            width, height = image.size
            metadata = {
                "format": "image", "extension": suffix[1:], "width": width,
                "height": height, "mode": image.mode, "inert_binary": True,
            }
        text = f"Image attachment metadata: {width} by {height} pixels, format {suffix[1:].upper()}."
    elif suffix in {".txt", ".py", ".sql"}:
        text = content.decode("utf-8-sig", errors="replace")
        metadata = {"format": suffix[1:], "inert_source": suffix in {".py", ".sql"}}
    elif suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        pages = [(page.extract_text() or "") for page in reader.pages[:100]]
        text = "\n\n".join(pages)
        metadata = {"format": "pdf", "pages": len(reader.pages), "pages_extracted": len(pages)}
    elif suffix == ".docx":
        from docx import Document

        document = Document(io.BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_rows = [
            " | ".join(_clean(cell.text) for cell in row.cells)
            for table in document.tables for row in table.rows
        ]
        text = "\n".join([*paragraphs, *table_rows])
        metadata = {
            "format": "docx", "paragraphs": len(document.paragraphs),
            "tables": len(document.tables), "table_rows": len(table_rows),
        }
    elif suffix == ".csv":
        frame = pd.read_csv(io.BytesIO(content))
        frame = frame.iloc[:, :100]
        rows = [" | ".join(_clean(value) for value in frame.columns)]
        rows.extend(" | ".join(_clean(value) for value in row) for row in frame.fillna("").itertuples(index=False, name=None))
        text = "\n".join(rows)
        metadata = {
            "format": "csv", "rows": len(frame), "columns": [str(value) for value in frame.columns],
            "sample_values": frame.head(5).fillna("").astype(str).to_dict(orient="records"),
        }
    else:
        from openpyxl import load_workbook

        row_counts: dict[str, int] = {}
        workbook_dimensions = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        try:
            row_counts = {
                str(sheet.title): max(0, int(sheet.max_row or 0) - 1)
                for sheet in workbook_dimensions.worksheets[:50]
            }
        finally:
            workbook_dimensions.close()
        workbook = pd.ExcelFile(io.BytesIO(content))
        sheet_details: list[dict[str, Any]] = []
        sections: list[str] = []
        for sheet_name in workbook.sheet_names[:50]:
            frame = workbook.parse(sheet_name=sheet_name, nrows=5000).iloc[:, :100]
            sheet_details.append({
                "name": str(sheet_name), "rows": row_counts.get(str(sheet_name), len(frame)),
                "rows_extracted": len(frame),
                "columns": [str(value) for value in frame.columns],
                "sample_values": frame.head(5).fillna("").astype(str).to_dict(orient="records"),
            })
            rows = [f"Sheet: {sheet_name}", " | ".join(_clean(value) for value in frame.columns)]
            rows.extend(" | ".join(_clean(value) for value in row) for row in frame.fillna("").itertuples(index=False, name=None))
            sections.append("\n".join(rows))
        workbook.close()
        text = "\n\n".join(sections)
        metadata = {"format": "xlsx", "sheet_names": workbook.sheet_names, "sheets": sheet_details}
    normalized = text.replace("\x00", "").strip()
    if not normalized:
        raise ValueError("No readable text was found in this file.")
    return normalized[:2_000_000], metadata


def chunk_text(text: str, size: int = 1400, overlap: int = 180) -> list[str]:
    chunks: list[str] = []
    cursor = 0
    while cursor < len(text) and len(chunks) < 1500:
        end = min(len(text), cursor + size)
        if end < len(text):
            boundary = max(text.rfind("\n", cursor + size // 2, end), text.rfind(". ", cursor + size // 2, end))
            if boundary > cursor:
                end = boundary + 1
        chunk = text[cursor:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        cursor = max(cursor + 1, end - overlap)
    return chunks
